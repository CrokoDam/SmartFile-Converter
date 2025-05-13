import os
import sys
import logging
from pathlib import Path
from typing import List, Optional
import typer
from PIL import Image
from PyPDF2 import PdfReader, PdfWriter
from pdf2docx import Converter as PDF2DOCXConverter
from docx import Document

app = typer.Typer()

# Setup logging
logging.basicConfig(
    filename='conversion.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Utility functions

def validate_file(file_path: Path):
    if not file_path.exists():
        logging.error(f"File not found: {file_path}")
        raise typer.BadParameter(f"File not found: {file_path}")
    return file_path

def extract_metadata(file_path: Path):
    metadata = {}
    if file_path.suffix.lower() == '.pdf':
        reader = PdfReader(str(file_path))
        info = reader.metadata
        metadata = {
            "Author": info.author,
            "Creator": info.creator,
            "Producer": info.producer,
            "Created": info.creation_date,
        }
    elif file_path.suffix.lower() == '.docx':
        doc = Document(str(file_path))
        props = doc.core_properties
        metadata = {
            "Author": props.author,
            "Created": props.created,
            "Title": props.title,
        }
    return metadata

def convert_docx_to_pdf(input_path: Path, output_path: Path):
    logging.warning("DOCX to PDF requires external tools not covered here.")
    typer.echo("DOCX to PDF conversion requires external tools like LibreOffice.")

def convert_pdf_to_docx(input_path: Path, output_path: Path):
    converter = PDF2DOCXConverter(str(input_path))
    converter.convert(str(output_path), start=0, end=None)
    converter.close()
    logging.info(f"Converted PDF to DOCX: {input_path} -> {output_path}")

def convert_pdf_to_text(input_path: Path, output_path: Path):
    reader = PdfReader(str(input_path))
    with open(output_path, 'w', encoding='utf-8') as f:
        for page in reader.pages:
            f.write(page.extract_text() + '\n')
    logging.info(f"Converted PDF to Text: {input_path} -> {output_path}")

def convert_text_to_pdf(input_path: Path, output_path: Path):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    with open(input_path, 'r', encoding='utf-8') as f:
        for line in f:
            pdf.cell(200, 10, txt=line.strip(), ln=True)
    pdf.output(str(output_path))
    logging.info(f"Converted Text to PDF: {input_path} -> {output_path}")

def convert_docx_to_markdown(input_path: Path, output_path: Path):
    doc = Document(str(input_path))
    with open(output_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            f.write(f"{para.text}\n\n")
    logging.info(f"Converted DOCX to Markdown: {input_path} -> {output_path}")

def convert_image_to_pdf(input_path: Path, output_path: Path):
    image = Image.open(input_path)
    image.convert('RGB').save(output_path)
    logging.info(f"Converted Image to PDF: {input_path} -> {output_path}")

def convert_pdf_to_image(input_path: Path, output_path: Path):
    typer.echo("PDF to Image conversion requires poppler or pdf2image.")
    logging.warning("PDF to Image not implemented.")

@app.command()
def convert(
    input: List[Path] = typer.Option(..., exists=True, help="Input file(s) to convert"),
    output: Optional[Path] = typer.Option(None, help="Output file or directory"),
    to: str = typer.Option(..., help="Target format: pdf, docx, txt, md, jpg, png"),
    metadata: bool = typer.Option(False, help="Extract metadata only")
):
    for file_path in input:
        file_path = validate_file(file_path)

        if metadata:
            data = extract_metadata(file_path)
            typer.echo(f"Metadata for {file_path.name}:")
            for key, value in data.items():
                typer.echo(f"  {key}: {value}")
            continue

        suffix = file_path.suffix.lower()
        output_file = output or Path(f"{file_path.stem}_converted.{to}")

        try:
            if suffix == '.pdf' and to == 'docx':
                convert_pdf_to_docx(file_path, output_file)
            elif suffix == '.pdf' and to == 'txt':
                convert_pdf_to_text(file_path, output_file)
            elif suffix == '.txt' and to == 'pdf':
                convert_text_to_pdf(file_path, output_file)
            elif suffix == '.docx' and to == 'md':
                convert_docx_to_markdown(file_path, output_file)
            elif suffix in ['.png', '.jpg', '.jpeg'] and to == 'pdf':
                convert_image_to_pdf(file_path, output_file)
            else:
                typer.echo(f"Conversion from {suffix} to {to} not supported.")
                logging.warning(f"Unsupported conversion: {suffix} to {to}")
        except Exception as e:
            typer.echo(f"Error converting {file_path.name}: {e}")
            logging.error(f"Error converting {file_path.name}: {e}")

if __name__ == "__main__":
    app()